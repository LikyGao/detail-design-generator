from __future__ import annotations

import base64
import io
import json
import re
from copy import deepcopy
from datetime import date
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, Twips

from tools.chapter_parser import _paragraph_structure
from tools.paragraph_numbering import calculate_paragraph_prefix


MIME_DOCX = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
PARAGRAPH_STYLES = {f"level_{level}" for level in range(7)}
# Derived from the standard template's numbering levels (twips), not a synthetic
# one-character-per-level ladder.  Explicit block/template values always win.
PARAGRAPH_FALLBACK_INDENTS = {0: 0, 1: 170, 2: 340, 3: 510, 4: 510, 5: 794, 6: 794}
PARAGRAPH_PREFIX_SEPARATOR = " "
# The parser's established native-level mapping is intentionally not numeric:
# Word/UI `スタイル0` (native ilvl 0) is represented as `level_1` internally.
UI_STYLE_ZERO_PARAGRAPH_STYLE = "level_1"
BODY_GROUP_SPACE_BEFORE = Pt(10.5)


def _replace_paragraph_text(paragraph, text: str) -> None:
    """Replace visible text while retaining the first run's formatting."""
    text = str(text or "")
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ""


def _replace_cell_text(cell, text: str) -> None:
    """Replace cell text while retaining the cell and first paragraph formatting."""
    text = str(text or "")
    if not cell.paragraphs:
        cell.add_paragraph(text)
        return
    first = cell.paragraphs[0]
    _replace_paragraph_text(first, text)
    for paragraph in cell.paragraphs[1:]:
        _replace_paragraph_text(paragraph, "")


def _set_update_fields(document: Document) -> None:
    settings = document.settings._element
    current = settings.find(qn("w:updateFields"))
    if current is None:
        current = OxmlElement("w:updateFields")
        settings.append(current)
    current.set(qn("w:val"), "true")


def _mark_all_fields_dirty(document: Document) -> None:
    """Mark template TOC fields and newly inserted caption fields for refresh."""
    root = document._element
    for field in root.xpath(".//w:fldSimple"):
        field.set(qn("w:dirty"), "true")
    for field_char in root.xpath(".//w:fldChar"):
        if field_char.get(qn("w:fldCharType")) == "begin":
            field_char.set(qn("w:dirty"), "true")


def _remove_existing_body(document: Document) -> None:
    """Remove template sample body beginning at the first Heading 1 paragraph."""
    start_element = None
    for paragraph in document.paragraphs:
        if paragraph.style and paragraph.style.name == "Heading 1":
            start_element = paragraph._element
            break
    if start_element is None:
        return

    body = document._element.body
    children = list(body)
    try:
        start_index = children.index(start_element)
    except ValueError:
        return

    for element in children[start_index:]:
        if element.tag == qn("w:sectPr"):
            continue
        body.remove(element)


def _fill_cover(
    document: Document,
    client_name: str,
    project_name: str,
    version: str,
    issue_date: str,
    file_name: str,
    project_no: str,
) -> None:
    title_paragraphs = [p for p in document.paragraphs if p.style and p.style.name == "Title"]
    if len(title_paragraphs) >= 2:
        _replace_paragraph_text(title_paragraphs[0], client_name)
        _replace_paragraph_text(title_paragraphs[1], project_name)

    if document.tables:
        property_table = document.tables[0]
        values = [version, issue_date, file_name, project_no]
        for row_index, value in enumerate(values):
            if value is None or row_index >= len(property_table.rows):
                continue
            cells = property_table.rows[row_index].cells
            if len(cells) >= 3:
                _replace_cell_text(cells[2], value)


def _fill_revision_history(document: Document, revisions: list[dict[str, Any]]) -> None:
    if len(document.tables) < 2:
        return
    table = document.tables[1]
    if len(table.rows) < 2:
        return

    base_row = deepcopy(table.rows[1]._tr)
    tbl = table._tbl
    while len(table.rows) > 1:
        tbl.remove(table.rows[-1]._tr)

    for item in revisions:
        tbl.append(deepcopy(base_row))
        row = table.rows[-1]
        values = [
            item.get("issue_date") or item.get("date") or "",
            item.get("version") or "",
            item.get("editor") or item.get("author") or "",
            item.get("description") or item.get("content") or "",
        ]
        for cell, value in zip(row.cells, values):
            _replace_cell_text(cell, value)


def _flatten_selected(nodes: Iterable[dict[str, Any]]) -> Iterable[dict[str, Any]]:
    for node in nodes or []:
        if not isinstance(node, dict):
            continue
        if node.get("selected", True):
            yield node
            yield from _flatten_selected(node.get("children") or [])


def _append_simple_field(paragraph, instruction: str, result_text: str) -> None:
    """Append a simple Word field and keep a visible fallback result."""
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), instruction)
    field.set(qn("w:dirty"), "true")

    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    text.text = str(result_text or "")
    run.append(text)
    field.append(run)
    paragraph._p.append(field)


def _caption_title(block: dict[str, Any], fallback_title: str) -> str:
    explicit = str(block.get("caption") or block.get("title") or "").strip()
    if explicit:
        return explicit

    file_name = str(block.get("fileName") or block.get("file_name") or "").strip()
    if file_name:
        stem = Path(file_name).stem.strip()
        if stem:
            return stem

    fallback = str(fallback_title or "").strip()
    return fallback or "タイトル未設定"


def _add_caption(
    document: Document,
    kind: str,
    caption: str,
    chapter_number: str,
    sequence_number: int,
) -> None:
    """Add a real Word caption using STYLEREF + SEQ fields.

    The template's table-of-figures fields can only discover native SEQ
    fields. Plain text styled as Caption is not enough.
    """
    paragraph = document.add_paragraph(style="Caption")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run(f"【{kind} ")
    _append_simple_field(paragraph, r"STYLEREF 1 \s", chapter_number)
    paragraph.add_run("-")
    _append_simple_field(
        paragraph,
        rf"SEQ {kind} \* ARABIC \s 1",
        str(sequence_number),
    )
    if caption:
        paragraph.add_run(f"：{caption}")
    paragraph.add_run("】")


def _add_table(document: Document, block: dict[str, Any]) -> None:
    headers = [str(v or "") for v in block.get("headers") or []]
    rows = block.get("rows") or []
    col_count = max(len(headers), max((len(r or []) for r in rows), default=0), 1)
    table = document.add_table(rows=1, cols=col_count)
    table.style = "Table Grid"

    if headers:
        for index in range(col_count):
            value = headers[index] if index < len(headers) else ""
            _replace_cell_text(table.rows[0].cells[index], value)
            for run in table.rows[0].cells[index].paragraphs[0].runs:
                run.bold = True
                run.font.name = "Meiryo UI"
    else:
        tbl = table._tbl
        tbl.remove(table.rows[0]._tr)

    for source_row in rows:
        target = table.add_row()
        source_row = source_row or []
        for index in range(col_count):
            value = source_row[index] if index < len(source_row) else ""
            _replace_cell_text(target.cells[index], value)


def _image_bytes(data_url: str) -> bytes:
    if not data_url:
        return b""
    if "," in data_url and data_url.startswith("data:"):
        data_url = data_url.split(",", 1)[1]
    return base64.b64decode(data_url)


def _add_figure(document: Document, block: dict[str, Any]) -> None:
    raw = _image_bytes(str(block.get("imageData") or block.get("image_data") or ""))
    if not raw:
        paragraph = document.add_paragraph("（画像未挿入）")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(io.BytesIO(raw), width=Inches(6.1))


def _collect_paragraph_prototypes(document: Document) -> dict[str, dict[str, Any]]:
    """Build canonical, template-native style descriptors before body removal.

    Descriptors are indexed independently by editor style, Word style id and
    Word style name. Paragraph-local properties from a template sample are not
    copied; fresh paragraphs inherit their layout from the native style and
    numbering definition just as they do when the style is applied in Word.
    """
    result: dict[str, dict[str, Any]] = {}
    for paragraph in document.paragraphs:
        if paragraph.style and paragraph.style.name in {"Heading 1", "Heading 2", "Heading 3"}:
            continue
        structure = _paragraph_structure(document, paragraph)
        paragraph_style = structure.get("paragraph_style")
        if paragraph_style not in PARAGRAPH_STYLES - {"level_0"} or structure.get("native_ilvl") is None:
            continue
        style = paragraph.style
        descriptor = {
            **structure,
            "style_id": str(style.style_id or ""),
            "style_name": str(style.name or ""),
            "canonical_pPr": deepcopy(style.element.pPr) if style.element.pPr is not None else None,
        }
        for key in (paragraph_style, f"id:{descriptor['style_id']}", f"name:{descriptor['style_name']}"):
            result.setdefault(key, descriptor)
    return result


def _descriptor_for(block: dict[str, Any], descriptors: dict[str, dict[str, Any]]):
    for key in (f"id:{block.get('word_style_id', '')}",
                f"name:{block.get('word_style_name', '')}",
                str(block.get("paragraph_style") or "")):
        if key and key in descriptors:
            candidate = descriptors[key]
            if isinstance(candidate, dict) and candidate.get("style_name"):
                return candidate
    return None


def _set_native_num_pr(paragraph, num_id: str, ilvl: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    old = p_pr.find(qn("w:numPr"))
    if old is not None:
        p_pr.remove(old)
    num_pr = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), str(ilvl))
    number = OxmlElement("w:numId")
    number.set(qn("w:val"), str(num_id))
    num_pr.extend((level, number))
    p_pr.append(num_pr)


def _numbering_instance_for_group(document, descriptor, group_id, instances, start=1):
    base_num_id = str(descriptor.get("num_id") or "")
    if not base_num_id or not group_id:
        return base_num_id
    try:
        start = int(start)
    except (TypeError, ValueError):
        start = 1
    key = (base_num_id, str(group_id))
    if key in instances:
        return instances[key]
    root = document.part.numbering_part.element
    nums = root.findall(qn("w:num"))
    source = next((n for n in nums if n.get(qn("w:numId")) == base_num_id), None)
    if source is None:
        return base_num_id
    new_id = str(max([int(n.get(qn("w:numId"), "0")) for n in nums] + [0]) + 1)
    clone = deepcopy(source)
    clone.set(qn("w:numId"), new_id)
    ilvl = int(descriptor.get("native_ilvl") or 0)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), str(ilvl))
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), str(start))
    override.append(start_override)
    clone.append(override)
    root.append(clone)
    instances[key] = new_id
    return new_id


def _native_numbering_restart(block: dict[str, Any], descriptor: dict[str, Any]) -> bool:
    """Return whether this paragraph needs a direct numbering instance.

    A native style already supplies its normal numbering and geometry through
    styles.xml and numbering.xml.  A paragraph-local numPr is only needed to
    select a cloned num instance when an actual numbered sequence has an
    independent group/restart.  Bullet and arrow groups do not have a counter
    to restart and must remain style-only.
    """
    if not block.get("list_group_id"):
        return False
    number_format = str(descriptor.get("number_format") or "").lower()
    return number_format not in {"", "bullet", "none"}


def _remove_numbering_properties(paragraph_properties) -> None:
    """Remove template numbering while retaining its other paragraph formatting."""
    if paragraph_properties is None:
        return
    num_properties = paragraph_properties.find(qn("w:numPr"))
    if num_properties is not None:
        paragraph_properties.remove(num_properties)


def _valid_indent_twips(value: Any) -> int | None:
    """Return a non-negative Word indent in twips, or None when unavailable."""
    if isinstance(value, bool) or value in (None, ""):
        return None
    try:
        twips = int(value)
    except (TypeError, ValueError):
        return None
    return twips if twips >= 0 else None


def _restore_body_indents(paragraph, block: dict[str, Any], paragraph_style: str) -> None:
    """Restore indents made ineffective when the template numPr is removed."""
    left_indent = _valid_indent_twips(block.get("left_indent_twips"))
    if left_indent is None:
        left_indent = PARAGRAPH_FALLBACK_INDENTS[int(paragraph_style[-1])]
    paragraph.paragraph_format.left_indent = Twips(left_indent)

    hanging_indent = _valid_indent_twips(block.get("hanging_indent_twips"))
    first_line_indent = _valid_indent_twips(block.get("first_line_indent_twips"))
    if hanging_indent is not None:
        paragraph.paragraph_format.first_line_indent = Twips(-hanging_indent)
    elif first_line_indent is not None:
        paragraph.paragraph_format.first_line_indent = Twips(first_line_indent)


def _is_ui_style_zero(block: dict[str, Any], descriptor: dict[str, Any] | None) -> bool:
    """Identify Word/UI スタイル0 through the parser's existing mapping."""
    if descriptor is not None:
        return (
            descriptor.get("paragraph_style") == UI_STYLE_ZERO_PARAGRAPH_STYLE
            and descriptor.get("native_ilvl") == 0
        )
    return (
        block.get("word_style_name") == "スタイル0"
        or block.get("paragraph_style") == UI_STYLE_ZERO_PARAGRAPH_STYLE
    )


def _apply_body_spacing(paragraph, block: dict[str, Any], descriptor: dict[str, Any] | None,
                        is_first_body_block: bool) -> None:
    """Separate later top-level body groups without creating empty paragraphs."""
    if _is_ui_style_zero(block, descriptor) and not is_first_body_block:
        paragraph.paragraph_format.space_before = BODY_GROUP_SPACE_BEFORE


def _add_body_paragraph(document: Document, block: dict[str, Any], prototypes: dict[str, Any],
                        counters: dict[int, int], numbering_instances: dict | None = None,
                        *, is_first_body_block: bool = True):
    numbering_instances = numbering_instances if numbering_instances is not None else {}
    paragraph_style = str(block.get("paragraph_style") or "level_0")
    if paragraph_style not in PARAGRAPH_STYLES:
        paragraph_style = "level_0"
    text = str(block.get("text") or "").replace("\r\n", "\n").replace("\r", "\n")

    descriptor = _descriptor_for(block, prototypes)
    if descriptor:
        word_style_name = descriptor["style_name"]
        try:
            document.styles[word_style_name]
        except KeyError:
            pass
        else:
            paragraph = document.add_paragraph(style=word_style_name)
            if _native_numbering_restart(block, descriptor):
                start = block.get("start_override")
                if start is None:
                    start = block.get("numbering_start")
                if start is None:
                    start = descriptor.get("start_override") or descriptor.get("numbering_start") or 1
                num_id = _numbering_instance_for_group(
                    document, descriptor, block.get("list_group_id"), numbering_instances, start)
                _set_native_num_pr(paragraph, num_id, int(descriptor.get("native_ilvl") or 0))
            paragraph.add_run(text)  # marker is rendered by Word, never literal text
            _apply_body_spacing(paragraph, block, descriptor, is_first_body_block)
            return paragraph

    # Compatibility for callers supplying a native style identity without a
    # descriptor map. The complete generator normally takes the descriptor path.
    word_style_name = str(block.get("word_style_name") or "")
    if word_style_name:
        try:
            document.styles[word_style_name]
        except KeyError:
            pass
        else:
            paragraph = document.add_paragraph(style=word_style_name)
            paragraph.add_run(text)
            _apply_body_spacing(paragraph, block, None, is_first_body_block)
            return paragraph

    paragraph = document.add_paragraph(style="Normal")
    prototype = prototypes.get(paragraph_style)
    if prototype is not None:
        if paragraph._p.pPr is not None:
            paragraph._p.remove(paragraph._p.pPr)
        paragraph_properties = deepcopy(prototype)
        _remove_numbering_properties(paragraph_properties)
        paragraph._p.insert(0, paragraph_properties)
        # A prototype can point at a list style whose numPr lives in styles.xml.
        # Keep the copied direct formatting, but use a non-list paragraph style so
        # literal prefixes are not combined with inherited automatic numbering.
        paragraph.style = "Normal"

    _restore_body_indents(paragraph, block, paragraph_style)
    prefix = calculate_paragraph_prefix(
        paragraph_style,
        counters,
        str(block.get("list_group_id") or "") or None,
        str(block.get("marker_type") or "") or None,
        block.get("start_override") if block.get("start_override") is not None
        else block.get("numbering_start"),
    )
    separator = PARAGRAPH_PREFIX_SEPARATOR if prefix and text else ""
    paragraph.add_run(f"{prefix}{separator}{text}")
    _apply_body_spacing(paragraph, block, descriptor, is_first_body_block)
    return paragraph


def _add_blocks(
    document: Document,
    blocks: list[dict[str, Any]],
    *,
    chapter_number: str,
    section_title: str,
    caption_counters: dict[str, int],
    paragraph_prototypes: dict[str, Any],
    numbering_scope: str = "section",
) -> None:
    paragraph_counters: dict = {}
    has_body_block = False
    blocks = _inherit_adjacent_list_groups(blocks or [])
    for block in blocks:
        if not isinstance(block, dict):
            continue
        block_type = str(block.get("type") or "paragraph")
        if block_type == "paragraph":
            scoped_block = block
            if str(block.get("paragraph_style") or "") == "level_1":
                scoped_block = dict(block)
                scoped_block["list_group_id"] = f"level-1:{numbering_scope}"
            _add_body_paragraph(document, scoped_block, paragraph_prototypes, paragraph_counters,
                                caption_counters.setdefault("_numbering_instances", {}),
                                is_first_body_block=not has_body_block)
            has_body_block = True
        elif block_type == "table":
            has_body_block = True
            caption_counters["表"] += 1
            _add_caption(
                document,
                "表",
                _caption_title(block, section_title),
                chapter_number,
                caption_counters["表"],
            )
            _add_table(document, block)
            document.add_paragraph("")
        elif block_type == "figure":
            has_body_block = True
            caption_counters["図"] += 1
            _add_caption(
                document,
                "図",
                _caption_title(block, section_title),
                chapter_number,
                caption_counters["図"],
            )
            _add_figure(document, block)
            document.add_paragraph("")
        elif block_type == "table_placeholder":
            has_body_block = True
            caption_counters["表"] += 1
            _add_caption(
                document,
                "表",
                _caption_title(block, section_title),
                chapter_number,
                caption_counters["表"],
            )
            paragraph = document.add_paragraph("（表は別途作成）")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif block_type == "figure_placeholder":
            has_body_block = True
            caption_counters["図"] += 1
            _add_caption(
                document,
                "図",
                _caption_title(block, section_title),
                chapter_number,
                caption_counters["図"],
            )
            paragraph = document.add_paragraph("（図は別途作成）")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _inherit_adjacent_list_groups(blocks: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """Attach newly inserted numbered paragraphs to an adjacent explicit group.

    Unknown JSON fields already pass through the tool unchanged.  This operates on
    a shallow copy, so persisted template metadata and caller input are untouched.
    """
    result = [dict(block) if isinstance(block, dict) else block for block in blocks]
    numbered = {"level_1", "level_4", "level_6"}
    for index, block in enumerate(result):
        if not isinstance(block, dict) or block.get("type", "paragraph") != "paragraph":
            continue
        style = str(block.get("paragraph_style") or "level_0")
        if style not in numbered or block.get("list_group_id"):
            continue
        adjacent_groups = []
        for direction in (-1, 1):
            cursor = index + direction
            while 0 <= cursor < len(result):
                candidate = result[cursor]
                if not isinstance(candidate, dict) or candidate.get("type", "paragraph") != "paragraph":
                    break
                candidate_style = str(candidate.get("paragraph_style") or "level_0")
                if candidate_style == style and candidate.get("list_group_id"):
                    adjacent_groups.append(str(candidate["list_group_id"]))
                    break
                # Descriptive bullets/children may sit inside one numbered item.
                if candidate_style in numbered:
                    break
                cursor += direction
        if adjacent_groups and len(set(adjacent_groups)) == 1:
            block["list_group_id"] = adjacent_groups[0]
    return result


def _chapter_number_from_node(node: dict[str, Any], fallback: int) -> str:
    display = str(node.get("display_number") or "").strip()
    if display:
        match = re.match(r"^([0-9０-９]+)", display)
        if match:
            return match.group(1)
    source = str(node.get("source_template_number") or "").strip()
    if source:
        match = re.match(r"^([0-9０-９]+)", source)
        if match:
            return match.group(1)
    return str(fallback)


def _add_chapters(document: Document, chapters: list[dict[str, Any]], paragraph_prototypes: dict[str, Any]) -> None:
    first_level_one = True
    chapter_index = 0
    current_chapter_number = "1"
    caption_counters = {"表": 0, "図": 0}

    for node_index, node in enumerate(_flatten_selected(chapters)):
        level = int(node.get("level") or 1)
        level = min(max(level, 1), 3)
        if level == 1:
            chapter_index += 1
            current_chapter_number = _chapter_number_from_node(node, chapter_index)
            caption_counters = {"表": 0, "図": 0}
            if not first_level_one:
                document.add_page_break()
            first_level_one = False

        title = str(node.get("title") or "(無題)")
        document.add_paragraph(title, style=f"Heading {level}")
        _add_blocks(
            document,
            node.get("blocks") or [],
            chapter_number=current_chapter_number,
            section_title=title,
            caption_counters=caption_counters,
            paragraph_prototypes=paragraph_prototypes,
            numbering_scope=str(node.get("id") or node_index),
        )


def parse_json_array(value: Any, field_name: str) -> list[dict[str, Any]]:
    if value in (None, ""):
        return []
    if isinstance(value, list):
        result = value
    else:
        try:
            result = json.loads(str(value))
        except json.JSONDecodeError as exc:
            raise ValueError(f"{field_name} is not valid JSON: {exc}") from exc
    if not isinstance(result, list):
        raise ValueError(f"{field_name} must be a JSON array")
    return result


def generate_standard_docx(
    template_source: bytes | bytearray | str | Path,
    *,
    client_name: str,
    project_name: str,
    version: str = "1.0",
    issue_date: str = "",
    project_no: str = "-",
    file_name: str = "基本設計書.docx",
    revisions: list[dict[str, Any]] | None = None,
    chapters: list[dict[str, Any]] | None = None,
) -> bytes:
    issue_date = issue_date or date.today().strftime("%Y/%m/%d")
    if isinstance(template_source, (bytes, bytearray)):
        document = Document(io.BytesIO(bytes(template_source)))
    else:
        document = Document(str(template_source))
    _fill_cover(document, client_name, project_name, version, issue_date, file_name, project_no)
    _fill_revision_history(document, revisions or [])
    paragraph_prototypes = _collect_paragraph_prototypes(document)
    _remove_existing_body(document)
    _add_chapters(document, chapters or [], paragraph_prototypes)
    _mark_all_fields_dirty(document)
    _set_update_fields(document)

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()
