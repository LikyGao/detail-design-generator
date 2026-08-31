from __future__ import annotations

import hashlib
import io
import json
import re
from datetime import datetime, timezone
from pathlib import Path
from typing import Any
from zipfile import BadZipFile, ZipFile

from docx import Document

from tools.chapter_parser import parse_template_chapters


DOCUMENT_TYPES = ("server_storage", "network", "cloud", "full")
DEFAULT_DOCUMENT_TYPE = "full"

# v0.0.7 registry. Each type has exactly one current template and version.
REGISTRY_KEY = "standard_word_template:v3:registry"
REGISTRY_SCHEMA_VERSION = 4

# v0.0.5-v0.0.6 single-template keys retained for compatibility fallback.
ACTIVE_TEMPLATE_KEY = "standard_word_template:current"
ACTIVE_META_KEY = "standard_word_template:active_meta"

# v0.0.2-v0.0.4 key retained only for compatibility cleanup helpers.
LEGACY_HISTORY_KEY = "standard_word_template:history"

MAX_TEMPLATE_BYTES = 50 * 1024 * 1024
# Keep each plugin KV backwards-invocation comfortably below gateway/body limits.
STORAGE_CHUNK_BYTES = 256 * 1024


def _json_bytes(value: Any) -> bytes:
    return json.dumps(value, ensure_ascii=False, separators=(",", ":")).encode("utf-8")


def _is_missing_storage_key_error(exc: Exception) -> bool:
    message = str(exc or "").strip().lower()
    markers = (
        "you have not set it",
        "key doesn't exist",
        "key does not exist",
        "key not found",
        "record not found",
        "no data found",
    )
    return any(marker in message for marker in markers)


def _get_optional(storage, key: str) -> bytes | None:
    try:
        return storage.get(key)
    except Exception as exc:
        if _is_missing_storage_key_error(exc):
            return None
        raise


def _delete_optional(storage, key: str) -> None:
    try:
        storage.delete(key)
    except Exception as exc:
        if not _is_missing_storage_key_error(exc):
            return


def _chunk_key(key: str, index: int) -> str:
    return f"{key}:part:{index:04d}"


def _set_payload(storage, key: str, payload: bytes) -> int:
    """Store bytes in one KV entry or deterministic chunks.

    Returns 0 for a single entry, otherwise the number of chunk entries.
    Older metadata without a chunk count remains readable.
    """
    if not key:
        raise ValueError("保存先キーが空です。")
    if len(payload) <= STORAGE_CHUNK_BYTES:
        try:
            storage.set(key, payload)
        except Exception as exc:
            raise RuntimeError(
                f"プラグインストレージへの保存に失敗しました: {key} "
                f"({len(payload)} bytes)"
            ) from exc
        return 0

    count = (len(payload) + STORAGE_CHUNK_BYTES - 1) // STORAGE_CHUNK_BYTES
    for index in range(count):
        start = index * STORAGE_CHUNK_BYTES
        chunk = payload[start : start + STORAGE_CHUNK_BYTES]
        chunk_key = _chunk_key(key, index)
        try:
            storage.set(chunk_key, chunk)
        except Exception as exc:
            raise RuntimeError(
                f"プラグインストレージへの分割保存に失敗しました: "
                f"{chunk_key} ({len(chunk)} bytes, {index + 1}/{count})"
            ) from exc
    return count


def _get_payload(storage, key: str, chunk_count: int = 0) -> bytes | None:
    if not key:
        return None
    count = int(chunk_count or 0)
    if count <= 0:
        return _get_optional(storage, key)

    parts: list[bytes] = []
    for index in range(count):
        chunk_key = _chunk_key(key, index)
        part = _get_optional(storage, chunk_key)
        if part is None:
            raise ValueError(
                f"分割保存データが不足しています: {chunk_key} ({index + 1}/{count})"
            )
        parts.append(part)
    return b"".join(parts)


def _read_json(
    storage, key: str, default: Any, chunk_count: int = 0
) -> Any:
    raw = _get_payload(storage, key, chunk_count)
    if not raw:
        return default
    try:
        return json.loads(raw.decode("utf-8"))
    except (UnicodeDecodeError, json.JSONDecodeError):
        return default


def normalize_document_type(value: Any, *, default_full: bool = False) -> str:
    document_type = str(value or "").strip().lower()
    if not document_type and default_full:
        document_type = DEFAULT_DOCUMENT_TYPE
    if document_type not in DOCUMENT_TYPES:
        allowed = ", ".join(DOCUMENT_TYPES)
        raise ValueError(f"document_typeは次のいずれかを指定してください: {allowed}")
    return document_type


def _safe_key_part(value: str) -> str:
    normalized = re.sub(r"[^a-zA-Z0-9_.-]+", "_", str(value or "").strip())
    return normalized.strip("._-") or "current"


def _storage_keys(document_type: str) -> dict[str, str]:
    safe_type = _safe_key_part(document_type)
    prefix = f"standard_word_template:v3:{safe_type}"
    return {
        "docx_key": f"{prefix}:docx",
        "master_key": f"{prefix}:master",
        "chapter_list_key": f"{prefix}:chapter_list",
        "section_contents_key": f"{prefix}:section_contents",
    }


def _empty_registry() -> dict[str, Any]:
    return {
        "schema_version": REGISTRY_SCHEMA_VERSION,
        "document_types": {document_type: None for document_type in DOCUMENT_TYPES},
        "updated_at": "",
    }


def _load_registry(storage) -> dict[str, Any]:
    registry = _read_json(storage, REGISTRY_KEY, None)
    if not isinstance(registry, dict):
        registry = _empty_registry()
    registry["schema_version"] = REGISTRY_SCHEMA_VERSION
    document_types = registry.get("document_types")
    if not isinstance(document_types, dict):
        document_types = {}
    for document_type in DOCUMENT_TYPES:
        document_types.setdefault(document_type, None)
    registry["document_types"] = document_types
    return registry


def _save_registry(storage, registry: dict[str, Any]) -> None:
    registry["updated_at"] = datetime.now(timezone.utc).isoformat()
    storage.set(REGISTRY_KEY, _json_bytes(registry))


def infer_template_version(filename: str) -> str:
    stem = Path(str(filename or "")).stem
    patterns = (
        r"(?:^|[_\-\s])v(?:er(?:sion)?)?[._\-\s]*([0-9]+(?:\.[0-9]+)*)$",
        r"(?:^|[_\-\s])([0-9]+(?:\.[0-9]+)+)$",
    )
    for pattern in patterns:
        match = re.search(pattern, stem, flags=re.IGNORECASE)
        if match:
            return match.group(1)
    return "unspecified"


def validate_template_bytes(template_bytes: bytes) -> dict[str, Any]:
    errors: list[str] = []
    warnings: list[str] = []

    if not template_bytes:
        return {"valid": False, "errors": ["テンプレートファイルが空です。"], "warnings": []}
    if len(template_bytes) > MAX_TEMPLATE_BYTES:
        errors.append(f"テンプレートは{MAX_TEMPLATE_BYTES // (1024 * 1024)}MB以下にしてください。")

    media_count = 0
    try:
        with ZipFile(io.BytesIO(template_bytes)) as archive:
            names = set(archive.namelist())
            for required in ("[Content_Types].xml", "word/document.xml", "word/styles.xml"):
                if required not in names:
                    errors.append(f"DOCX内部に必要ファイルがありません: {required}")
            media_count = len(
                [name for name in names if name.startswith("word/media/") and not name.endswith("/")]
            )
    except BadZipFile:
        errors.append("有効なDOCX（ZIP形式）ではありません。")

    document = None
    if not errors:
        try:
            document = Document(io.BytesIO(template_bytes))
        except Exception as exc:
            errors.append(f"Wordテンプレートを開けません: {exc}")

    summary: dict[str, Any] = {
        "file_size": len(template_bytes),
        "media_count": media_count,
        "paragraph_count": 0,
        "table_count": 0,
        "section_count": 0,
    }

    if document is not None:
        summary.update(
            {
                "paragraph_count": len(document.paragraphs),
                "table_count": len(document.tables),
                "section_count": len(document.sections),
            }
        )

        style_names = {style.name for style in document.styles}
        required_styles = {"Normal", "Heading 1", "Heading 2", "Heading 3", "Caption"}
        missing_styles = sorted(required_styles - style_names)
        if missing_styles:
            errors.append("必要なWordスタイルがありません: " + ", ".join(missing_styles))

        title_paragraphs = [
            p for p in document.paragraphs if p.style is not None and p.style.name == "Title"
        ]
        if len(title_paragraphs) < 2:
            errors.append("表紙の顧客名・プロジェクト名に使用するTitle段落が2つ必要です。")

        heading_one = [
            p for p in document.paragraphs if p.style is not None and p.style.name == "Heading 1"
        ]
        if not heading_one:
            errors.append("本文挿入位置を判定するHeading 1段落がありません。")

        if len(document.tables) < 2:
            errors.append("Document Property表と改訂履歴表の2表が必要です。")
        else:
            property_table = document.tables[0]
            revision_table = document.tables[1]
            if len(property_table.rows) < 4 or any(
                len(row.cells) < 3 for row in property_table.rows[:4]
            ):
                errors.append("先頭のDocument Property表は4行・3列以上必要です。")
            if len(revision_table.rows) < 2 or len(revision_table.rows[0].cells) < 4:
                errors.append("2番目の改訂履歴表は見出し＋雛形行、4列以上必要です。")

        if media_count == 0:
            warnings.append("word/media内に画像がありません。Logo等が意図どおりか確認してください。")

    return {
        "valid": not errors,
        "errors": errors,
        "warnings": warnings,
        "summary": summary,
    }


def register_typed_template(
    storage,
    *,
    template_bytes: bytes,
    filename: str,
    document_type: str,
    template_version: str,
) -> dict[str, Any]:
    document_type = normalize_document_type(document_type, default_full=True)
    filename = filename or "基本設計書_template.docx"
    template_version = str(template_version or "").strip() or infer_template_version(filename)

    validation = validate_template_bytes(template_bytes)
    if not validation["valid"]:
        raise ValueError("テンプレート検証に失敗しました: " + " / ".join(validation["errors"]))

    parsed = parse_template_chapters(template_bytes)
    keys = _storage_keys(document_type)
    digest = hashlib.sha256(template_bytes).hexdigest()
    now = datetime.now(timezone.utc).isoformat()

    metadata = {
        "id": digest[:16],
        "sha256": digest,
        "filename": filename,
        "document_type": document_type,
        "template_version": template_version,
        "updated_at": now,
        "size": len(template_bytes),
        "validation": validation,
        "chapter_summary": parsed["summary"],
        "section_content_summary": {
            "section_count": parsed["summary"]["total_count"],
            "sections_with_text": parsed["summary"]["sections_with_text"],
            "body_paragraph_count": parsed["summary"]["body_paragraph_count"],
            "body_character_count": parsed["summary"]["body_character_count"],
            "skipped_body_paragraph_count": parsed["summary"]["skipped_body_paragraph_count"],
        },
        "parse_warnings": parsed["warnings"],
        "source": "persistent_storage",
        "mode": "one_current_per_type",
        **keys,
    }

    # Write payloads first and commit the registry last. Large payloads are split
    # because Dify Cloud transports storage operations through backwards invocation.
    docx_chunk_count = _set_payload(storage, keys["docx_key"], template_bytes)
    master_chunk_count = _set_payload(
        storage, keys["master_key"], _json_bytes(parsed["master_json"])
    )
    chapter_list_chunk_count = _set_payload(
        storage, keys["chapter_list_key"], _json_bytes(parsed["chapter_list_json"])
    )
    section_contents_chunk_count = _set_payload(
        storage,
        keys["section_contents_key"],
        _json_bytes(parsed["section_contents_json"]),
    )
    metadata.update(
        {
            "storage_chunk_bytes": STORAGE_CHUNK_BYTES,
            "docx_chunk_count": docx_chunk_count,
            "master_chunk_count": master_chunk_count,
            "chapter_list_chunk_count": chapter_list_chunk_count,
            "section_contents_chunk_count": section_contents_chunk_count,
        }
    )

    registry = _load_registry(storage)
    registry["document_types"][document_type] = metadata
    _save_registry(storage, registry)

    return {
        "metadata": metadata,
        "master_json": parsed["master_json"],
        "chapter_list_json": parsed["chapter_list_json"],
        "section_content_summary": metadata["section_content_summary"],
        "warnings": parsed["warnings"],
    }


def register_template(
    storage,
    *,
    template_bytes: bytes,
    filename: str,
    document_type: str = DEFAULT_DOCUMENT_TYPE,
    template_version: str = "",
) -> dict[str, Any]:
    """Compatibility registration entry point.

    Calls from v0.0.6 workflows only provide template_file; those calls now register the
    template as document_type=full while preserving the existing tool name.
    """
    result = register_typed_template(
        storage,
        template_bytes=template_bytes,
        filename=filename,
        document_type=document_type or DEFAULT_DOCUMENT_TYPE,
        template_version=template_version,
    )
    return result["metadata"]


def _migrate_legacy_current_to_full(storage) -> dict[str, Any] | None:
    """Lazily copy the v0.0.6 single current template into the v0.0.7 full slot.

    The legacy keys are intentionally retained, so an interrupted migration does not break
    the existing default workflow.
    """
    registry = _load_registry(storage)
    current = registry["document_types"].get(DEFAULT_DOCUMENT_TYPE)
    if isinstance(current, dict):
        return current

    legacy_bytes = _get_optional(storage, ACTIVE_TEMPLATE_KEY)
    if not legacy_bytes:
        return None
    legacy_meta = _read_json(storage, ACTIVE_META_KEY, {})
    filename = (
        str(legacy_meta.get("filename") or "基本設計書_template.docx")
        if isinstance(legacy_meta, dict)
        else "基本設計書_template.docx"
    )
    version = (
        str(legacy_meta.get("template_version") or "").strip()
        if isinstance(legacy_meta, dict)
        else ""
    ) or infer_template_version(filename)
    result = register_typed_template(
        storage,
        template_bytes=legacy_bytes,
        filename=filename,
        document_type=DEFAULT_DOCUMENT_TYPE,
        template_version=version,
    )
    return result["metadata"]


def get_registered_template_metadata(storage, document_type: str) -> dict[str, Any] | None:
    document_type = normalize_document_type(document_type)
    registry = _load_registry(storage)
    metadata = registry["document_types"].get(document_type)
    if isinstance(metadata, dict):
        return metadata
    if document_type == DEFAULT_DOCUMENT_TYPE:
        return _migrate_legacy_current_to_full(storage)
    return None


def get_registered_template(
    storage,
    *,
    document_type: str,
    template_version: str = "",
    template_id: str = "",
) -> tuple[bytes, dict[str, Any]]:
    document_type = normalize_document_type(document_type)
    metadata = get_registered_template_metadata(storage, document_type)
    if metadata is None:
        raise ValueError(f"document_type={document_type} の標準テンプレートは未登録です。")

    requested_id = str(template_id or "").strip()
    current_id = str(metadata.get("id") or "")
    if requested_id and requested_id != current_id:
        raise ValueError(
            f"document_type={document_type} の現在テンプレートIDは {current_id or '未設定'} です。"
            f"指定ID {requested_id} は登録されていません。"
        )

    requested_version = str(template_version or "").strip()
    current_version = str(metadata.get("template_version") or "")
    if requested_version and requested_version != current_version:
        raise ValueError(
            f"document_type={document_type} の現在版は {current_version or '未設定'} です。"
            f"指定版 {requested_version} は登録されていません。"
        )

    template_bytes = _get_payload(
        storage,
        str(metadata.get("docx_key") or ""),
        int(metadata.get("docx_chunk_count") or 0),
    )
    if not template_bytes:
        raise ValueError(f"document_type={document_type} のDOCXデータが見つかりません。")
    return template_bytes, metadata


def get_registered_master(storage, document_type: str) -> dict[str, Any]:
    document_type = normalize_document_type(document_type)
    metadata = get_registered_template_metadata(storage, document_type)
    if metadata is None:
        raise ValueError(f"document_type={document_type} の標準テンプレートは未登録です。")

    master = _read_json(
        storage,
        str(metadata.get("master_key") or ""),
        None,
        int(metadata.get("master_chunk_count") or 0),
    )
    chapter_list = _read_json(
        storage,
        str(metadata.get("chapter_list_key") or ""),
        None,
        int(metadata.get("chapter_list_chunk_count") or 0),
    )
    if not isinstance(master, list) or not isinstance(chapter_list, list):
        raise ValueError(f"document_type={document_type} の章節Masterデータが見つかりません。")
    return {
        "document_type": document_type,
        "template_version": metadata.get("template_version") or "",
        "template": metadata,
        "master_json": master,
        "chapter_list_json": chapter_list,
    }


def get_registered_section_contents(
    storage,
    *,
    document_type: str,
    template_version: str = "",
    template_id: str = "",
) -> dict[str, Any]:
    document_type = normalize_document_type(document_type)
    metadata = get_registered_template_metadata(storage, document_type)
    if metadata is None:
        raise ValueError(f"document_type={document_type} の標準テンプレートは未登録です。")

    requested_id = str(template_id or "").strip()
    current_id = str(metadata.get("id") or "")
    if requested_id and requested_id != current_id:
        raise ValueError(
            f"document_type={document_type} の現在テンプレートIDは {current_id or '未設定'} です。"
            f"指定ID {requested_id} は登録されていません。"
        )

    requested_version = str(template_version or "").strip()
    current_version = str(metadata.get("template_version") or "")
    if requested_version and requested_version != current_version:
        raise ValueError(
            f"document_type={document_type} の現在版は {current_version or '未設定'} です。"
            f"指定版 {requested_version} は登録されていません。"
        )

    section_contents_key = str(metadata.get("section_contents_key") or "")
    section_contents = (
        _read_json(
            storage,
            section_contents_key,
            None,
            int(metadata.get("section_contents_chunk_count") or 0),
        )
        if section_contents_key
        else None
    )
    if not isinstance(section_contents, list):
        raise ValueError(
            f"document_type={document_type} の章節標準本文データがありません。"
            "v0.0.8以降のプラグインで標準テンプレートを再登録してください。"
        )
    return {
        "document_type": document_type,
        "template_id": current_id,
        "template_version": current_version,
        "template": metadata,
        "section_contents": section_contents,
    }


def _get_legacy_or_bundled_template(
    storage, bundled_template_path: str | Path
) -> tuple[bytes, dict[str, Any]]:
    template_bytes = _get_optional(storage, ACTIVE_TEMPLATE_KEY)
    metadata = _read_json(storage, ACTIVE_META_KEY, None)

    if template_bytes:
        if not isinstance(metadata, dict):
            metadata = {
                "filename": "基本設計書_template.docx",
                "updated_at": "",
                "size": len(template_bytes),
                "source": "persistent_storage",
                "mode": "legacy_single_template",
            }
        return template_bytes, metadata

    if isinstance(metadata, dict) and metadata.get("storage_key"):
        legacy_bytes = _get_optional(storage, str(metadata["storage_key"]))
        if legacy_bytes:
            public_meta = {k: v for k, v in metadata.items() if k != "storage_key"}
            public_meta["mode"] = "legacy_active"
            return legacy_bytes, public_meta

    path = Path(bundled_template_path)
    bundled_bytes = path.read_bytes()
    validation = validate_template_bytes(bundled_bytes)
    bundled_meta = {
        "id": "bundled-default",
        "sha256": hashlib.sha256(bundled_bytes).hexdigest(),
        "filename": path.name,
        "document_type": DEFAULT_DOCUMENT_TYPE,
        "template_version": "bundled",
        "updated_at": "",
        "size": len(bundled_bytes),
        "validation": validation,
        "source": "bundled_default",
        "mode": "fallback",
    }
    return bundled_bytes, bundled_meta


def select_template(
    storage,
    bundled_template_path: str | Path,
    *,
    document_type: str = "",
    template_version: str = "",
    template_id: str = "",
) -> tuple[bytes, dict[str, Any]]:
    requested_type = str(document_type or "").strip().lower()
    if requested_type:
        return get_registered_template(
            storage,
            document_type=requested_type,
            template_version=template_version,
            template_id=template_id,
        )

    # Backward-compatible default: full -> legacy v0.0.6 current -> bundled template.
    full_metadata = get_registered_template_metadata(storage, DEFAULT_DOCUMENT_TYPE)
    if full_metadata is not None:
        return get_registered_template(
            storage,
            document_type=DEFAULT_DOCUMENT_TYPE,
            template_version=template_version,
            template_id=template_id,
        )
    if template_version or template_id:
        raise ValueError("template_versionまたはtemplate_idを指定する場合はdocument_typeも指定してください。")
    return _get_legacy_or_bundled_template(storage, bundled_template_path)


def get_active_template(storage, bundled_template_path: str | Path) -> tuple[bytes, dict[str, Any]]:
    """v0.0.6-compatible default template selector."""
    return select_template(storage, bundled_template_path)


def list_registered_templates(storage) -> dict[str, Any]:
    registry = _load_registry(storage)
    templates: dict[str, Any] = {}
    for document_type in DOCUMENT_TYPES:
        metadata = registry["document_types"].get(document_type)
        templates[document_type] = metadata if isinstance(metadata, dict) else None
    return {
        "schema_version": REGISTRY_SCHEMA_VERSION,
        "templates": templates,
        "registered_count": sum(1 for value in templates.values() if value is not None),
    }


def get_template_status(storage, bundled_template_path: str | Path) -> dict[str, Any]:
    _, active = select_template(storage, bundled_template_path)
    registered = list_registered_templates(storage)
    return {
        # Retained for existing workflows that read active_template.
        "active_template": active,
        "templates": registered["templates"],
        "registered_count": registered["registered_count"],
        "document_types": list(DOCUMENT_TYPES),
        "history_enabled": False,
        "storage_mode": "one_current_per_type",
        "default_selection_order": ["full", "legacy_current", "bundled_default"],
    }
