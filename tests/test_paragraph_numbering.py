import io
import importlib.util
import sys
import unittest
from pathlib import Path


ROOT = Path(__file__).resolve().parents[1]
PLUGIN_ROOT = ROOT / "plugins" / "personal" / "standard_word_generator"
sys.path.insert(0, str(PLUGIN_ROOT))

from tools.paragraph_numbering import calculate_paragraph_prefix


class ParagraphNumberingTest(unittest.TestCase):
    def prefixes(self, styles):
        counters = {}
        return [calculate_paragraph_prefix(style, counters) for style in styles]

    def test_nested_numbering_and_continuation(self):
        self.assertEqual(
            self.prefixes(
                ["level_1", "level_2", "level_3", "level_4", "level_5", "level_4", "level_1"]
            ),
            ["（1）", "・", "➢", "①", "・", "②", "（2）"],
        )

    def test_parent_restarts_deeper_numbering(self):
        self.assertEqual(
            self.prefixes(["level_1", "level_4", "level_4", "level_1", "level_4"]),
            ["（1）", "①", "②", "（2）", "①"],
        )

    def test_two_explicit_circle_groups_do_not_share_counters(self):
        counters = {}
        items = [("group-a", "level_4")] * 2 + [("group-b", "level_4")] * 3
        self.assertEqual(
            [calculate_paragraph_prefix(style, counters, group) for group, style in items],
            ["①", "②", "①", "②", "③"],
        )

    def test_explicit_group_honors_numbering_start(self):
        counters = {}
        self.assertEqual(
            [calculate_paragraph_prefix("level_4", counters, "restart", "circle", 3)
             for _ in range(2)],
            ["③", "④"],
        )

    def test_native_six_level_numbering_and_style_inheritance(self):
        if importlib.util.find_spec("docx") is None:
            self.skipTest("python-docx is not installed")
        from docx import Document
        from docx.enum.style import WD_STYLE_TYPE
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from tools.chapter_parser import _paragraph_structure

        document = Document()
        root = document.part.numbering_part.element
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), "42")
        formats = ["decimal", "bullet", "decimalEnclosedCircle", "bullet",
                   "bullet", "decimalEnclosedCircle"]
        markers = ["(%1)", "・", "%3", "\uf0d8", "・", "%6"]
        style_ids = []
        for ilvl, (number_format, marker) in enumerate(zip(formats, markers)):
            style = document.styles.add_style(f"Native level {ilvl}", WD_STYLE_TYPE.PARAGRAPH)
            style_ids.append(style.style_id)
            level = OxmlElement("w:lvl")
            level.set(qn("w:ilvl"), str(ilvl))
            for tag, value in (("w:start", "1"), ("w:numFmt", number_format),
                               ("w:lvlText", marker), ("w:pStyle", style.style_id)):
                child = OxmlElement(tag)
                child.set(qn("w:val"), value)
                level.append(child)
            p_pr = OxmlElement("w:pPr")
            ind = OxmlElement("w:ind")
            ind.set(qn("w:left"), str(240 * (ilvl + 1)))
            ind.set(qn("w:hanging"), "120")
            p_pr.append(ind)
            level.append(p_pr)
            if ilvl == 3:
                r_pr = OxmlElement("w:rPr")
                fonts = OxmlElement("w:rFonts")
                fonts.set(qn("w:ascii"), "Wingdings")
                fonts.set(qn("w:hAnsi"), "Wingdings")
                r_pr.append(fonts)
                level.append(r_pr)
            abstract.append(level)
        root.append(abstract)
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), "42")
        abstract_ref = OxmlElement("w:abstractNumId")
        abstract_ref.set(qn("w:val"), "42")
        num.append(abstract_ref)
        root.append(num)

        # The concrete style has no numPr; its basedOn ancestor is linked by
        # numbering-level pStyle, which must still resolve the native ilvl.
        inherited = document.styles.add_style("Inherited native level 3", WD_STYLE_TYPE.PARAGRAPH)
        inherited.base_style = document.styles[style_ids[3]]

        structures = []
        for ilvl, style_id in enumerate(style_ids):
            selected_style = inherited if ilvl == 3 else document.styles[style_id]
            paragraph = document.add_paragraph(f"body {ilvl}", style=selected_style)
            structures.append(_paragraph_structure(document, paragraph, paragraph.text))

        self.assertEqual([item["list_level"] for item in structures], list(range(6)))
        self.assertEqual([item["paragraph_style"] for item in structures],
                         ["level_1", "level_2", "level_4", "level_3", "level_5", "level_6"])
        self.assertEqual(
            [calculate_paragraph_prefix(item["paragraph_style"], {}) for item in structures],
            ["（1）", "・", "①", "➢", "・", "①"],
        )
        self.assertEqual([item["number_format"] for item in structures], formats)
        self.assertEqual(structures[3]["symbol_font"], "Wingdings")
        self.assertEqual(structures[3]["marker_type"], "arrow")
        self.assertEqual(structures[1]["paragraph_style"], "level_2")
        self.assertEqual(structures[4]["paragraph_style"], "level_5")
        self.assertNotEqual(structures[1]["paragraph_style"], structures[4]["paragraph_style"])
        self.assertEqual(structures[0]["left_indent_twips"], 240)
        self.assertEqual(structures[5]["hanging_indent_twips"], 120)
        self.assertEqual(structures[2]["native_ilvl"], 2)
        self.assertEqual(structures[2]["word_style_id"], style_ids[2])
        self.assertEqual(structures[2]["word_style_name"], "Native level 2")
        self.assertEqual(structures[3]["word_style_based_on_id"], style_ids[3])
        self.assertEqual(structures[2]["numbering_p_style"], style_ids[2])
        self.assertEqual(structures[2]["abstract_num_id"], "42")

    def test_current_style_exact_match_precedes_based_on_and_falls_back_only_when_absent(self):
        if importlib.util.find_spec("docx") is None:
            self.skipTest("python-docx is not installed")
        from docx import Document
        from docx.enum.style import WD_STYLE_TYPE
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from tools.chapter_parser import _paragraph_structure

        document = Document()
        base = document.styles.add_style("スタイル3", WD_STYLE_TYPE.PARAGRAPH)
        current = document.styles.add_style("スタイル4", WD_STYLE_TYPE.PARAGRAPH)
        current.base_style = base
        undefined = document.styles.add_style("スタイル5", WD_STYLE_TYPE.PARAGRAPH)
        undefined.base_style = current

        root = document.part.numbering_part.element
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), "88")
        # Put the ancestor first to reproduce the v0.0.14 unordered-set failure.
        for ilvl, style_id, marker in ((3, base.style_id, "\uf0d8"),
                                       (4, current.style_id, "・")):
            level = OxmlElement("w:lvl")
            level.set(qn("w:ilvl"), str(ilvl))
            for tag, value in (("w:numFmt", "bullet"), ("w:lvlText", marker),
                               ("w:pStyle", style_id)):
                child = OxmlElement(tag)
                child.set(qn("w:val"), value)
                level.append(child)
            if ilvl == 3:
                r_pr = OxmlElement("w:rPr")
                fonts = OxmlElement("w:rFonts")
                fonts.set(qn("w:ascii"), "Wingdings")
                r_pr.append(fonts)
                level.append(r_pr)
            abstract.append(level)
        root.append(abstract)
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), "88")
        ref = OxmlElement("w:abstractNumId")
        ref.set(qn("w:val"), "88")
        num.append(ref)
        root.append(num)

        own = document.add_paragraph("own", style=current)
        inherited = document.add_paragraph("inherited", style=undefined)
        own_structure = _paragraph_structure(document, own, own.text)
        inherited_structure = _paragraph_structure(document, inherited, inherited.text)

        self.assertEqual(own_structure["native_ilvl"], 4)
        self.assertEqual(own_structure["paragraph_style"], "level_5")
        self.assertEqual(own_structure["marker_type"], "bullet")
        self.assertEqual(inherited_structure["native_ilvl"], 4)
        self.assertEqual(inherited_structure["paragraph_style"], "level_5")
        self.assertEqual(inherited_structure["word_style_based_on_id"], current.style_id)

    def test_group_insert_and_delete_renumber_locally(self):
        counters = {}
        self.assertEqual(
            [calculate_paragraph_prefix("level_4", counters, "same") for _ in range(3)],
            ["①", "②", "③"],
        )
        counters = {}
        self.assertEqual(
            [calculate_paragraph_prefix("level_4", counters, "same") for _ in range(2)],
            ["①", "②"],
        )

    def test_literal_bullet_is_classified_before_it_is_stripped(self):
        if importlib.util.find_spec("docx") is None:
            self.skipTest("python-docx is not installed")
        from docx import Document
        from tools.chapter_parser import _paragraph_structure, _strip_literal_paragraph_marker

        document = Document()
        paragraph = document.add_paragraph("・B")
        structure = _paragraph_structure(document, paragraph, paragraph.text)
        self.assertEqual(structure["marker_type"], "bullet")
        self.assertEqual(structure["paragraph_style"], "level_2")
        self.assertEqual(_strip_literal_paragraph_marker(paragraph.text, structure["paragraph_style"]), "B")

    def test_bullet_circle_regression_visible_structure_and_indent(self):
        if importlib.util.find_spec("docx") is None:
            self.skipTest("python-docx is not installed")
        from docx import Document
        from tools.docx_builder import _add_body_paragraph

        document = Document()
        counters = {}
        blocks = [
            {"paragraph_style": "level_4", "marker_type": "circle", "list_group_id": "g", "text": "A", "left_indent_twips": 510},
            {"paragraph_style": "level_5", "marker_type": "bullet", "text": "B", "left_indent_twips": 794},
            {"paragraph_style": "level_6", "marker_type": "circle", "list_group_id": "nested", "text": "C", "left_indent_twips": 794},
        ]
        for block in blocks:
            _add_body_paragraph(document, block, {}, counters)
        self.assertEqual([p.text for p in document.paragraphs], ["① A", "・ B", "① C"])
        self.assertEqual([p.paragraph_format.left_indent.twips for p in document.paragraphs], [510, 794, 794])

    def test_body_paragraph_uses_existing_native_word_style_without_literal_prefix(self):
        if importlib.util.find_spec("docx") is None:
            self.skipTest("python-docx is not installed")
        from docx import Document
        from docx.enum.style import WD_STYLE_TYPE
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        from tools.docx_builder import _add_body_paragraph

        document = Document()
        native_style = document.styles.add_style("スタイル4", WD_STYLE_TYPE.PARAGRAPH)
        num_properties = OxmlElement("w:numPr")
        level = OxmlElement("w:ilvl")
        level.set(qn("w:val"), "3")
        num_id = OxmlElement("w:numId")
        num_id.set(qn("w:val"), "1")
        num_properties.extend((level, num_id))
        native_style.element.get_or_add_pPr().append(num_properties)

        _add_body_paragraph(
            document,
            {
                "paragraph_style": "level_4",
                "word_style_name": "スタイル4",
                "text": "OS",
            },
            {},
            {},
        )

        paragraph = document.paragraphs[0]
        self.assertEqual(paragraph.style.name, "スタイル4")
        self.assertEqual(paragraph.text, "OS")
        self.assertNotIn("①", paragraph.text)
        self.assertIsNotNone(native_style.element.pPr.numPr)

    def test_body_paragraph_falls_back_when_word_style_name_is_missing(self):
        if importlib.util.find_spec("docx") is None:
            self.skipTest("python-docx is not installed")
        from docx import Document

        from tools.docx_builder import _add_body_paragraph

        document = Document()
        _add_body_paragraph(
            document,
            {
                "paragraph_style": "level_4",
                "word_style_name": "存在しないスタイル",
                "text": "OS",
            },
            {},
            {},
        )

        paragraph = document.paragraphs[0]
        self.assertEqual(paragraph.style.name, "Normal")
        self.assertEqual(paragraph.text, "① OS")

    def test_template_output_uses_literal_prefixes_without_effective_numbering(self):
        if importlib.util.find_spec("docx") is None:
            self.skipTest("python-docx is not installed")
        from docx import Document

        from tools.chapter_parser import _numbering_values
        from tools.docx_builder import generate_standard_docx

        template = PLUGIN_ROOT / "templates" / "基本設計書_template.docx"
        blocks = [
            {"type": "paragraph", "paragraph_style": style, "text": text}
            for style, text in [
                ("level_1", "バックアップ"),
                ("level_4", "OS"),
                ("level_4", "データ"),
                ("level_1", "リストア"),
                ("level_4", "設定"),
            ]
        ]
        output = generate_standard_docx(
            template,
            client_name="test",
            project_name="test",
            chapters=[{"level": 1, "title": "test", "blocks": blocks}],
        )

        expected = ["（1） バックアップ", "① OS", "② データ", "（2） リストア", "① 設定"]
        paragraphs = {
            paragraph.text: paragraph
            for paragraph in Document(io.BytesIO(output)).paragraphs
            if paragraph.text in expected
        }

        self.assertEqual(list(paragraphs), expected)
        for paragraph in paragraphs.values():
            self.assertIsNone(_numbering_values(paragraph))

    def test_body_paragraph_fallback_indents_increase_by_level(self):
        if importlib.util.find_spec("docx") is None:
            self.skipTest("python-docx is not installed")
        from docx import Document

        from tools.chapter_parser import _numbering_values
        from tools.docx_builder import _add_body_paragraph, PARAGRAPH_FALLBACK_INDENTS

        document = Document()
        counters = {}
        for level in range(7):
            _add_body_paragraph(
                document,
                {"paragraph_style": f"level_{level}", "text": f"level {level}"},
                {},
                counters,
            )

        indents = [paragraph.paragraph_format.left_indent.twips for paragraph in document.paragraphs]
        self.assertEqual(indents, list(PARAGRAPH_FALLBACK_INDENTS.values()))
        self.assertEqual(indents, sorted(indents))
        self.assertLessEqual(max(b - a for a, b in zip(indents, indents[1:])), 284)
        for paragraph in document.paragraphs:
            self.assertIsNone(_numbering_values(paragraph))

    def test_body_paragraph_prefers_original_indents_and_removes_prototype_numbering(self):
        if importlib.util.find_spec("docx") is None:
            self.skipTest("python-docx is not installed")
        from copy import deepcopy

        from docx import Document
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        from tools.chapter_parser import _numbering_values
        from tools.docx_builder import _add_body_paragraph

        document = Document()
        source = document.add_paragraph(style="Normal")
        num_properties = OxmlElement("w:numPr")
        num_id = OxmlElement("w:numId")
        num_id.set(qn("w:val"), "1")
        num_properties.append(num_id)
        source._p.get_or_add_pPr().append(num_properties)
        prototype = deepcopy(source._p.pPr)
        source._element.getparent().remove(source._element)

        _add_body_paragraph(
            document,
            {
                "paragraph_style": "level_2",
                "text": "original indent",
                "left_indent_twips": 975,
                "first_line_indent_twips": 120,
            },
            {"level_2": prototype},
            {},
        )
        _add_body_paragraph(
            document,
            {
                "paragraph_style": "level_3",
                "text": "hanging indent",
                "left_indent_twips": 1230,
                "first_line_indent_twips": 90,
                "hanging_indent_twips": 240,
            },
            {"level_3": prototype},
            {},
        )

        first, hanging = document.paragraphs
        self.assertEqual(first.paragraph_format.left_indent.twips, 975)
        self.assertEqual(first.paragraph_format.first_line_indent.twips, 120)
        self.assertEqual(hanging.paragraph_format.left_indent.twips, 1230)
        self.assertEqual(hanging.paragraph_format.first_line_indent.twips, -240)
        self.assertIsNone(_numbering_values(first))
        self.assertIsNone(_numbering_values(hanging))


class NativeCanonicalIndentGenerationTest(unittest.TestCase):
    def test_native_numbering_materializes_distinct_canonical_indents(self):
        if importlib.util.find_spec("docx") is None:
            self.skipTest("python-docx is not installed")
        from docx import Document
        from docx.enum.style import WD_STYLE_TYPE
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from tools.chapter_parser import _numbering_values
        from tools.docx_builder import _add_body_paragraph, _collect_paragraph_prototypes

        document = Document()
        root = document.part.numbering_part.element
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), "123")
        styles = {}
        canonical = {1: (360, 120), 2: (720, 180), 4: (1320, 240), 5: (1740, 300)}
        formats = {1: ("bullet", "・"), 2: ("decimalEnclosedCircle", "%3"),
                   4: ("bullet", "・"), 5: ("decimalEnclosedCircle", "%6")}
        for ilvl in range(6):
            style = document.styles.add_style(f"Canonical native {ilvl}", WD_STYLE_TYPE.PARAGRAPH)
            styles[ilvl] = style
            level = OxmlElement("w:lvl")
            level.set(qn("w:ilvl"), str(ilvl))
            number_format, marker = formats.get(ilvl, ("decimal", f"%{ilvl + 1}"))
            for tag, value in (("w:start", "1"), ("w:numFmt", number_format),
                               ("w:lvlText", marker), ("w:pStyle", style.style_id)):
                child = OxmlElement(tag); child.set(qn("w:val"), value); level.append(child)
            if ilvl in canonical:
                left, hanging = canonical[ilvl]
                p_pr = OxmlElement("w:pPr")
                indent = OxmlElement("w:ind")
                indent.set(qn("w:left"), str(left)); indent.set(qn("w:hanging"), str(hanging))
                p_pr.append(indent); level.append(p_pr)
            abstract.append(level)
        root.append(abstract)
        num = OxmlElement("w:num"); num.set(qn("w:numId"), "123")
        ref = OxmlElement("w:abstractNumId"); ref.set(qn("w:val"), "123")
        num.append(ref); root.append(num)

        for ilvl in (1, 2, 4, 5):
            paragraph = document.add_paragraph(f"prototype {ilvl}", style=styles[ilvl])
            num_pr = OxmlElement("w:numPr")
            level = OxmlElement("w:ilvl"); level.set(qn("w:val"), str(ilvl))
            number = OxmlElement("w:numId"); number.set(qn("w:val"), "123")
            num_pr.extend((level, number)); paragraph._p.get_or_add_pPr().append(num_pr)

        descriptors = _collect_paragraph_prototypes(document)
        for paragraph in list(document.paragraphs):
            paragraph._element.getparent().remove(paragraph._element)
        blocks = [
            {"id": "shallow-bullet", "type": "paragraph", "paragraph_style": "level_2", "text": "shallow bullet"},
            {"id": "deep-bullet", "type": "paragraph", "paragraph_style": "level_5", "text": "deep bullet"},
            {"id": "shallow-circle", "type": "paragraph", "paragraph_style": "level_4", "text": "shallow circle", "list_group_id": "circle-a"},
            {"id": "deep-circle", "type": "paragraph", "paragraph_style": "level_6", "text": "deep circle", "list_group_id": "circle-b"},
        ]
        for block in blocks:
            _add_body_paragraph(document, block, descriptors, {}, {})

        generated = document.paragraphs
        self.assertEqual([p.text for p in generated], [b["text"] for b in blocks])
        self.assertEqual([_numbering_values(p) for p in generated],
                         [("123", 1), ("123", 4), ("124", 2), ("125", 5)])
        self.assertEqual([p.style.name for p in generated],
                         [styles[i].name for i in (1, 4, 2, 5)])
        self.assertEqual([p.paragraph_format.left_indent.twips for p in generated],
                         [canonical[i][0] for i in (1, 4, 2, 5)])
        self.assertEqual([p.paragraph_format.first_line_indent.twips for p in generated],
                         [-canonical[i][1] for i in (1, 4, 2, 5)])
        self.assertGreater(generated[1].paragraph_format.left_indent.twips,
                           generated[0].paragraph_format.left_indent.twips)
        self.assertGreater(generated[3].paragraph_format.left_indent.twips,
                           generated[2].paragraph_format.left_indent.twips)
        for paragraph in generated:
            self.assertNotRegex(paragraph.text, r"^(?:・|[①-⑳]|➢|（\\d+）)")


if __name__ == "__main__":
    unittest.main()
