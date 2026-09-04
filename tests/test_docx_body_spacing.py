import io
import sys
import unittest
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn
from lxml import etree


ROOT = Path(__file__).resolve().parents[1]
PLUGIN_ROOT = ROOT / "plugins" / "personal" / "standard_word_generator"
sys.path.insert(0, str(PLUGIN_ROOT))

from tools.docx_builder import BODY_GROUP_SPACE_BEFORE, generate_standard_docx


TEMPLATE = PLUGIN_ROOT / "templates" / "基本設計書_template.docx"


def paragraph(style, text):
    return {"type": "paragraph", "paragraph_style": style, "text": text}


def generated_document(blocks, *, level=2, title="Spacing regression"):
    output = generate_standard_docx(
        TEMPLATE,
        client_name="Client",
        project_name="Project",
        chapters=[{"id": "spacing", "level": level, "title": title, "blocks": blocks}],
    )
    return output, Document(io.BytesIO(output))


def by_text(document, text):
    return next(item for item in document.paragraphs if item.text == text)


class DocxBodySpacingRegressionTest(unittest.TestCase):
    def assert_no_direct_spacing(self, paragraph):
        spacing = paragraph._p.get_or_add_pPr().find(qn("w:spacing"))
        self.assertTrue(
            spacing is None
            or (spacing.get(qn("w:before")) is None and spacing.get(qn("w:after")) is None)
        )

    def test_later_style_zero_group_gets_spacing_but_first_and_children_do_not(self):
        output, document = generated_document([
            paragraph("level_1", "style zero A"),
            paragraph("level_2", "child A-1"),
            paragraph("level_2", "child A-2"),
            paragraph("level_1", "style zero B"),
            paragraph("level_2", "child B-1"),
        ])

        self.assert_no_direct_spacing(by_text(document, "style zero A"))
        self.assert_no_direct_spacing(by_text(document, "child A-1"))
        self.assert_no_direct_spacing(by_text(document, "child A-2"))
        self.assertEqual(by_text(document, "style zero B").paragraph_format.space_before,
                         BODY_GROUP_SPACE_BEFORE)
        self.assert_no_direct_spacing(by_text(document, "child B-1"))
        self.assertTrue(output)

    def test_every_child_style_remains_without_direct_before_or_after_spacing(self):
        _, document = generated_document([
            paragraph("level_1", "top"),
            paragraph("level_2", "style one"),
            paragraph("level_4", "style two"),
            paragraph("level_3", "style three"),
            paragraph("level_5", "style four"),
        ])

        for text in ("style one", "style two", "style three", "style four"):
            self.assert_no_direct_spacing(by_text(document, text))

    def test_new_style_zero_is_spaced_even_when_previous_style_is_style_four(self):
        _, document = generated_document([
            paragraph("level_1", "group A"),
            paragraph("level_2", "style one child"),
            paragraph("level_5", "style four child"),
            paragraph("level_1", "group B"),
        ])

        self.assertEqual(by_text(document, "group B").paragraph_format.space_before,
                         BODY_GROUP_SPACE_BEFORE)

    def test_first_style_zero_after_heading_two_has_no_group_spacing(self):
        _, document = generated_document(
            [paragraph("level_1", "first after heading")], level=2, title="Heading two"
        )

        self.assert_no_direct_spacing(by_text(document, "first after heading"))

    def test_body_to_next_heading_does_not_insert_an_empty_xml_paragraph(self):
        output = generate_standard_docx(
            TEMPLATE,
            client_name="Client",
            project_name="Project",
            chapters=[
                {"id": "one", "level": 2, "title": "First heading",
                 "blocks": [paragraph("level_1", "first body")]},
                {"id": "two", "level": 2, "title": "Second heading", "blocks": []},
            ],
        )
        with ZipFile(io.BytesIO(output)) as archive:
            root = etree.fromstring(archive.read("word/document.xml"))
        paragraphs = root.xpath("//w:body/w:p", namespaces=root.nsmap)
        first_body = next(i for i, item in enumerate(paragraphs)
                          if "".join(item.xpath(".//w:t/text()", namespaces=root.nsmap)) == "first body")
        second_heading = next(i for i, item in enumerate(paragraphs)
                              if "".join(item.xpath(".//w:t/text()", namespaces=root.nsmap)) == "Second heading")
        self.assertEqual(second_heading, first_body + 1)
        first_heading = next(i for i, item in enumerate(paragraphs)
                             if "".join(item.xpath(".//w:t/text()", namespaces=root.nsmap)) == "First heading")
        self.assertFalse(any(len(item) == 0 for item in paragraphs[first_heading:]))

    def test_spacing_survives_save_and_reopen_and_is_stored_in_document_xml(self):
        output, document = generated_document([
            paragraph("level_1", "persist A"),
            paragraph("level_2", "persist child"),
            paragraph("level_1", "persist B"),
        ])
        reopened = Document(io.BytesIO(output))
        target = by_text(reopened, "persist B")
        self.assertEqual(target.paragraph_format.space_before, BODY_GROUP_SPACE_BEFORE)
        spacing = target._p.get_or_add_pPr().find(qn("w:spacing"))
        self.assertEqual(spacing.get(qn("w:before")), "210")
        self.assertEqual(len(document.paragraphs), len(reopened.paragraphs))


if __name__ == "__main__":
    unittest.main()
